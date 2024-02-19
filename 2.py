import docx
import MTK2
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX


doc = docx.Document('C:/D/5 курс/progs/2 семестр/Сафарьян/2/1.docx')

if __name__ == '__main__':
    text = "Чести без труда не сыскать."

    TextMTK2 = MTK2.MTK2_code(text)
    print(TextMTK2)

    LenParagraphs = []
    OT = ""
    for paragraph in doc.paragraphs:
        string = ""
        for run in paragraph.runs:
            for char in run.text:
                string += char
        OT += string
        LenParagraphs.append(len(string))

    # print(OT)
    print(LenParagraphs)

    doc.paragraphs.clear()
    id_symb = 0
    for string_id in range(len(doc.paragraphs)):
        doc.paragraphs[string_id].clear()
        for id_rans in range(LenParagraphs[string_id]):
            run = doc.paragraphs[string_id].add_run(OT[id_symb])
            run.font.name = 'Helvetica'
            run.font.size = Pt(13.5)
            if id_symb < len(TextMTK2):
                if TextMTK2[id_symb] == '1':
                    run.font.color.rgb = RGBColor(0, 0, 1)
                else:
                    run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                run.font.color.rgb = RGBColor(0, 0, 0)
            id_symb += 1

    doc.save('C:/D/5 курс/progs/2 семестр/Сафарьян/2/01variant.docx')
    print(f"Документ был изменён и сохранён!")